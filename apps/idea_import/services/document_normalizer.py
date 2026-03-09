"""
DocumentNormalizerService — Dokument-Formate → plain text.

Unterstützte Formate: txt, md, docx, pdf, freetext
Portiert aus bfagent (ADR-081).
"""
from __future__ import annotations

import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)


class UnsupportedFormatError(Exception):
    """Nicht unterstütztes Dateiformat."""


class DocumentNormalizerService:
    SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}

    @classmethod
    def normalize_upload(cls, file_content: bytes, filename: str, encoding: str = "utf-8") -> tuple[str, str]:
        ext = Path(filename).suffix.lower()
        if ext == ".txt":
            return cls._decode(file_content, encoding), "txt"
        elif ext == ".md":
            return cls._strip_markdown(cls._decode(file_content, encoding)), "md"
        elif ext == ".docx":
            return cls._normalize_docx(file_content), "docx"
        elif ext == ".pdf":
            return cls._normalize_pdf(file_content), "pdf"
        raise UnsupportedFormatError(f"Format '{ext}' nicht unterstützt. Erlaubt: txt, md, docx, pdf")

    @classmethod
    def normalize_freetext(cls, text: str) -> tuple[str, str]:
        return cls._clean_whitespace(text), "freetext"

    @classmethod
    def _normalize_docx(cls, content: bytes) -> str:
        try:
            import io
            import docx
        except ImportError as exc:
            raise ImportError("python-docx nicht installiert: pip install python-docx") from exc
        doc = docx.Document(io.BytesIO(content))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_texts = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_texts:
                    parts.append(" | ".join(row_texts))
        return cls._clean_whitespace("\n\n".join(parts))

    @classmethod
    def _normalize_pdf(cls, content: bytes) -> str:
        try:
            import io
            from pypdf import PdfReader
        except ImportError as exc:
            raise ImportError("pypdf nicht installiert: pip install pypdf") from exc
        reader = PdfReader(io.BytesIO(content))
        pages = [p.extract_text().strip() for p in reader.pages if p.extract_text()]
        return cls._clean_whitespace("\n\n".join(pages))

    @staticmethod
    def _decode(content: bytes, encoding: str) -> str:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            return content.decode("latin-1", errors="replace")

    @staticmethod
    def _strip_markdown(text: str) -> str:
        text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", text)
        text = re.sub(r"_{1,3}([^_]+)_{1,3}", r"\1", text)
        text = re.sub(r"```[\s\S]*?```", "", text)
        text = re.sub(r"`[^`]+`", "", text)
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
        text = re.sub(r"!\[[^\]]*\]\([^)]+\)", "", text)
        return DocumentNormalizerService._clean_whitespace(text)

    @staticmethod
    def _clean_whitespace(text: str) -> str:
        text = re.sub(r"\r\n", "\n", text)
        text = re.sub(r"\r", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]+$", "", text, flags=re.MULTILINE)
        return text.strip()
